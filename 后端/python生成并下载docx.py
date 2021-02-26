# pip install python-docx
from docx import Document
from io import BytesIO


def get(self, request):
    response = HttpResponse(content_type='application/msword;charset=UTF-8')
    response['Content-Disposition'] = 'attachment;filename=境外收款银行账户信息.docx,'
    doc = Document()
    doc.add_heading('境外收款银行账户信息', level=1).paragraph_format.alignment = 1
    doc.add_paragraph('').paragraph_format.alignment = 1

    doc.add_paragraph('*收款人名称： ULILI')
    doc.add_paragraph('*收款人地址： UNIT NO.9 7 /F FOCAL IND CTRBLK A NO.21 MAN LOK ST HUNG HOM')
    doc.add_paragraph('*收款人账号： 382-562-1012457-01')
    doc.add_paragraph('*收款人开户行名称： BANK OF COMMUNICATIONS(HONG KONG)LTD')
    doc.add_paragraph('*收款人开户行地址： G/F,SHOP 1-3,NO.22-28 MODY ROAD,TSIM SHA TSUI,KOWLOON,HONG KONG')
    doc.add_paragraph('*SWIFT代码/人行行号： COMMHKHK')
    doc.add_paragraph('*银行号码： 989584002708 ')
    doc.add_paragraph('*收款人常驻国家（地区）名称： HONG KONG')
    doc.add_paragraph('*汇款附言（只限140个字位）： 其他私人旅游服务团费')

    file_io = BytesIO()
    doc.save(file_io)
    file_io.seek(0)
    response.write(file_io.getvalue())
    return response